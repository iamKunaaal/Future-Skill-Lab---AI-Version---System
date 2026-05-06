"""Shared DOCX helpers — used by exports.py (Phase 1) and materials_exports.py (Phase 2)."""
import re

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm


# ── Brand palette ───────────────────────────────────────────────────────────
PRIMARY      = RGBColor(0x63, 0x0E, 0xD4)   # purple
INK          = RGBColor(0x1F, 0x29, 0x37)   # near-black
MUTED        = RGBColor(0x64, 0x74, 0x8B)   # slate-500
ACCENT_TEAL  = RGBColor(0x0D, 0x94, 0x88)
ACCENT_AMBER = RGBColor(0xD9, 0x77, 0x06)
ACCENT_PINK  = RGBColor(0xEC, 0x48, 0x99)
ACCENT_BLUE  = RGBColor(0x25, 0x63, 0xEB)
RULE_GREY    = RGBColor(0xCB, 0xD5, 0xE1)


def set_cell_shading(cell, hex_color: str):
    """Apply background color to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def add_horizontal_rule(doc, color: str = 'CBD5E1'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_styled_run(paragraph, text, *, bold=False, italic=False,
                   size=11, color=INK, font='Calibri'):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return run


# ── Inline markdown parser (bold + italic + code) ─────────────────────────
_INLINE_PATTERN = re.compile(r'(\*\*[^*\n]+?\*\*|\*[^*\n]+?\*|`[^`\n]+?`)')


def render_inline(paragraph, text, *, base_size=11, base_color=INK):
    """Render bold/italic/code spans within a single line."""
    if not text:
        return
    parts = _INLINE_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            add_styled_run(paragraph, part[2:-2], bold=True,
                           size=base_size, color=base_color)
        elif part.startswith('*') and part.endswith('*'):
            add_styled_run(paragraph, part[1:-1], italic=True,
                           size=base_size, color=base_color)
        elif part.startswith('`') and part.endswith('`'):
            add_styled_run(paragraph, part[1:-1],
                           size=base_size - 1, color=PRIMARY,
                           font='Consolas')
        else:
            add_styled_run(paragraph, part, size=base_size, color=base_color)


# ── Markdown table helpers ────────────────────────────────────────────────
_TABLE_ROW_RE = re.compile(r'^\s*\|.*\|\s*$')
_TABLE_SEP_RE = re.compile(r'^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$')


def _split_md_row(row: str):
    s = row.strip()
    if s.startswith('|'):
        s = s[1:]
    if s.endswith('|'):
        s = s[:-1]
    return [c.strip() for c in s.split('|')]


def render_md_table(doc, raw_rows):
    """Render a markdown table block (list of pipe-separated rows) as a docx table."""
    rows = [_split_md_row(r) for r in raw_rows if r.strip()]
    rows = [r for r in rows if not all(re.fullmatch(r':?-{2,}:?', c or '') for c in r)]
    if not rows:
        return
    cols = max(len(r) for r in rows)
    rows = [r + [''] * (cols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Light Grid Accent 1'
    table.autofit = True

    header_cells = table.rows[0].cells
    for idx, text in enumerate(rows[0]):
        set_cell_shading(header_cells[idx], 'EDE9FE')
        cell_p = header_cells[idx].paragraphs[0]
        cell_p.paragraph_format.space_after = Pt(0)
        render_inline(cell_p, text, base_size=10, base_color=PRIMARY)
        for run in cell_p.runs:
            run.bold = True

    for r_idx, row in enumerate(rows[1:], start=1):
        for c_idx, text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell_p = cell.paragraphs[0]
            cell_p.paragraph_format.space_after = Pt(0)
            render_inline(cell_p, text, base_size=10, base_color=INK)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


# ── Markdown block parser ─────────────────────────────────────────────────
def render_markdown(doc, text: str):
    """Convert AI-generated markdown into well-formatted docx paragraphs."""
    if not text:
        return
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    lines = text.split('\n')

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped or re.fullmatch(r'-{3,}|_{3,}|\*{3,}', stripped):
            i += 1
            continue

        # Markdown table
        if _TABLE_ROW_RE.match(line) and '|' in stripped[1:]:
            tbl_rows = []
            while i < len(lines):
                nxt = lines[i].rstrip()
                if not nxt.strip():
                    break
                if _TABLE_ROW_RE.match(nxt) or _TABLE_SEP_RE.match(nxt):
                    tbl_rows.append(nxt)
                    i += 1
                    continue
                break
            data_rows = [r for r in tbl_rows if not _TABLE_SEP_RE.match(r)]
            if len(data_rows) >= 2:
                render_md_table(doc, tbl_rows)
                continue
            for tr in tbl_rows:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                render_inline(p, tr.strip())
            continue

        # Headings
        m = re.match(r'^(#{2,6})\s+(.+)$', stripped)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10 if level >= 4 else 14)
            p.paragraph_format.space_after = Pt(4)
            size_map = {2: 15, 3: 13, 4: 12, 5: 11, 6: 11}
            color = PRIMARY if level == 2 else (ACCENT_AMBER if level == 3 else INK)
            render_inline(p, heading_text, base_size=size_map.get(level, 12),
                          base_color=color)
            for run in p.runs:
                run.bold = True
            i += 1
            continue

        # Bullet list
        if re.match(r'^[\-\*•]\s+', stripped):
            while i < len(lines) and re.match(r'^[\-\*•]\s+', lines[i].strip()):
                item_text = re.sub(r'^[\-\*•]\s+', '', lines[i].strip())
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(2)
                render_inline(p, item_text)
                i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\.\s+', stripped):
            while i < len(lines) and re.match(r'^\d+\.\s+', lines[i].strip()):
                item_text = re.sub(r'^\d+\.\s+', '', lines[i].strip())
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.space_after = Pt(2)
                render_inline(p, item_text)
                i += 1
            continue

        # Blockquote
        if stripped.startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(re.sub(r'^>\s?', '', lines[i].strip()))
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(6)
            render_inline(p, ' '.join(quote_lines), base_color=MUTED)
            for run in p.runs:
                run.italic = True
            continue

        # Plain paragraph (collect contiguous lines)
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                break
            if re.match(r'^(#{2,6}\s|[\-\*•]\s|\d+\.\s|>)', nxt):
                break
            if _TABLE_ROW_RE.match(lines[i]) or _TABLE_SEP_RE.match(lines[i]):
                break
            if re.fullmatch(r'-{3,}|_{3,}|\*{3,}', nxt):
                break
            para_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        render_inline(p, ' '.join(para_lines))
