"""DOCX export for a Project — pulls latest content from DB so edits/regenerations are reflected."""
import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

from framework.models import Week, Session
from .models import Project, SessionContent
from .docx_utils import (
    PRIMARY, INK, MUTED, ACCENT_TEAL, ACCENT_AMBER,
    set_cell_shading      as _set_cell_shading,
    add_horizontal_rule   as _add_horizontal_rule,
    add_styled_run        as _add_styled_run,
    render_inline         as _render_inline,
    render_markdown       as _render_markdown,
    render_md_table       as _render_md_table,
)


# ── Cover & info blocks ───────────────────────────────────────────────────
def _add_cover(doc, project: Project):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(60)
    title.paragraph_format.space_after = Pt(4)
    _add_styled_run(title, 'NEORISE FSL · CURRICULUM EXPORT',
                    bold=True, size=10, color=PRIMARY)

    h = doc.add_paragraph()
    h.paragraph_format.space_after = Pt(20)
    _add_styled_run(h, project.topic, bold=True, size=26, color=INK)

    table = doc.add_table(rows=0, cols=2)
    table.autofit = True
    meta_rows = [
        ('Grade',          project.grade),
        ('Subject Track',  project.get_subject_track_display()),
        ('Tech Focus',     ', '.join(project.tech_competency) or '—'),
        ('Status',         project.get_status_display()),
        ('Created',        project.created_at.strftime('%d %b %Y')),
        ('Last Updated',   project.updated_at.strftime('%d %b %Y, %H:%M')),
        ('Exported On',    datetime.now().strftime('%d %b %Y, %H:%M')),
        ('Sessions',       f'{project.sessions_generated}/{project.sessions_total} generated · {project.sessions_approved} approved'),
    ]
    for label, value in meta_rows:
        row = table.add_row().cells
        row[0].width = Cm(4.5)
        _set_cell_shading(row[0], 'F1F5F9')
        p1 = row[0].paragraphs[0]
        _add_styled_run(p1, label.upper(), bold=True, size=9, color=MUTED)
        p2 = row[1].paragraphs[0]
        _add_styled_run(p2, str(value), size=10, color=INK)

    if project.description:
        doc.add_paragraph().paragraph_format.space_after = Pt(8)
        sub = doc.add_paragraph()
        _add_styled_run(sub, 'PROJECT NOTES', bold=True, size=9, color=MUTED)
        sub.paragraph_format.space_after = Pt(2)
        p = doc.add_paragraph()
        _add_styled_run(p, project.description, size=11, color=INK)

    doc.add_page_break()


def _add_week_section(doc, week: Week, project: Project, contents: dict):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(2)
    _add_styled_run(h, f'WEEK {week.number}', bold=True, size=10, color=MUTED)

    h2 = doc.add_paragraph()
    h2.paragraph_format.space_after = Pt(6)
    _add_styled_run(h2, week.phase, bold=True, size=20, color=PRIMARY)

    if week.description:
        sub = doc.add_paragraph()
        sub.paragraph_format.space_after = Pt(8)
        _add_styled_run(sub, week.description, italic=True, size=11, color=MUTED)

    _add_horizontal_rule(doc)

    first_session = week.sessions.order_by('number').first()
    weekly_brief = ''
    if first_session:
        sc = contents.get(first_session.number)
        if sc and sc.weekly_brief:
            weekly_brief = sc.weekly_brief

    if weekly_brief:
        bh = doc.add_paragraph()
        bh.paragraph_format.space_before = Pt(8)
        bh.paragraph_format.space_after = Pt(4)
        _add_styled_run(bh, 'WEEKLY BRIEF', bold=True, size=10, color=ACCENT_TEAL)
        _render_markdown(doc, weekly_brief)

    if week.kaushal_bodh_questions:
        kh = doc.add_paragraph()
        kh.paragraph_format.space_before = Pt(10)
        kh.paragraph_format.space_after = Pt(4)
        _add_styled_run(kh, 'KAUSHAL BODH — REFLECTION QUESTIONS',
                        bold=True, size=10, color=ACCENT_TEAL)
        for q in week.kaushal_bodh_questions:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            _add_styled_run(p, q, size=11, color=INK)


def _resolve_competencies(session: Session, project: Project):
    """Replace tech-slot placeholders with the project's actual tech competency selections."""
    tech_codes = project.tech_competency or []
    tech_descriptions = Project.TECH_DESCRIPTIONS

    def _sp_name(code):
        for c, label in Project.TECH_CHOICES:
            if c == code:
                return label.split(':')[0].strip()
        return code

    seen = set()
    out = []
    for comp in session.competencies.all():
        if comp.is_tech_slot:
            for code in tech_codes:
                if code in seen:
                    continue
                seen.add(code)
                out.append({
                    'msp_code': code,
                    'sp_name': _sp_name(code),
                    'description': tech_descriptions.get(code, ''),
                })
        else:
            if comp.msp_code in seen:
                continue
            seen.add(comp.msp_code)
            out.append({
                'msp_code': comp.msp_code,
                'sp_name': comp.sp_name,
                'description': comp.description,
            })
    return out


def _add_session(doc, session: Session, content: SessionContent | None, project: Project):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(2)
    _add_styled_run(h, f'SESSION {session.number} · BLOCK PERIOD {session.number}',
                    bold=True, size=9, color=MUTED)

    h2 = doc.add_paragraph()
    h2.paragraph_format.space_after = Pt(6)
    _add_styled_run(h2, session.name, bold=True, size=16, color=INK)

    status_text = '○ Not generated'
    status_color = MUTED
    if content and content.is_approved:
        status_text = '✓ Approved'
        status_color = ACCENT_TEAL
    elif content and content.ai_description:
        status_text = '◐ Pending review'
        status_color = ACCENT_AMBER
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)
    _add_styled_run(sp, status_text, bold=True, size=9, color=status_color)

    comps = _resolve_competencies(session, project)
    if comps:
        comp_table = doc.add_table(rows=1, cols=2)
        comp_table.autofit = True
        head = comp_table.rows[0].cells
        _set_cell_shading(head[0], 'F1F5F9')
        _set_cell_shading(head[1], 'F1F5F9')
        _add_styled_run(head[0].paragraphs[0], 'CODE',
                        bold=True, size=8, color=MUTED)
        _add_styled_run(head[1].paragraphs[0], 'COMPETENCY',
                        bold=True, size=8, color=MUTED)
        for c in comps:
            row = comp_table.add_row().cells
            _add_styled_run(row[0].paragraphs[0], c['msp_code'],
                            bold=True, size=10, color=PRIMARY,
                            font='Consolas')
            cell_p = row[1].paragraphs[0]
            _add_styled_run(cell_p, c['sp_name'], bold=True,
                            size=10, color=INK)
            desc_p = row[1].add_paragraph()
            _add_styled_run(desc_p, c['description'], size=9, color=MUTED)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    sh = doc.add_paragraph()
    sh.paragraph_format.space_before = Pt(6)
    sh.paragraph_format.space_after = Pt(4)
    _add_styled_run(sh, 'SESSION BREAKDOWN', bold=True, size=10, color=PRIMARY)

    if content and content.ai_description:
        _render_markdown(doc, content.ai_description)
    else:
        p = doc.add_paragraph()
        _add_styled_run(p, 'No AI-generated content yet for this session.',
                        italic=True, size=10, color=MUTED)


# ── Main entry point ─────────────────────────────────────────────────────
def build_project_docx(project: Project) -> io.BytesIO:
    """Build a formatted DOCX for a project with the latest DB content."""
    doc = Document()

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK

    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    _add_cover(doc, project)

    contents = {
        sc.session.number: sc
        for sc in project.session_contents.select_related('session__week').all()
    }

    weeks = (Week.objects
             .prefetch_related('sessions__competencies')
             .order_by('number'))

    for week in weeks:
        _add_week_section(doc, week, project, contents)
        for session in week.sessions.order_by('number'):
            _add_session(doc, session, contents.get(session.number), project)
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
